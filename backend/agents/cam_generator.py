"""
Agent 4 — CAM Generator Agent
Generates a professional Credit Appraisal Memo as a .docx Word document
using python-docx, covering all mandatory sections.
"""
from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from utils.prompts import CAM_EXECUTIVE_SUMMARY_PROMPT


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _log(agent: str, message: str, level: str = "INFO") -> Dict[str, Any]:
    return {
        "timestamp": datetime.now().strftime("%H:%M:%S"),
        "agent": agent,
        "message": message,
        "level": level,
    }


def _call_gemini_for_summary(state: Dict[str, Any]) -> str:
    """Ask Gemini to write the Executive Summary section."""
    scores = state.get("five_cs_scores", {})
    research = state.get("research_findings", {})
    extracted = state.get("extracted_financials", {})
    rec = state.get("final_recommendation", {})

    five_cs_summary = "\n".join(
        f"  {c.upper()}: {scores.get(c, {}).get('score', 0)}/100"
        for c in ["character", "capacity", "capital", "collateral", "conditions"]
    )
    red_flags = extracted.get("red_flags", [])
    key_findings = [
        f"[{f.get('severity')}] {f.get('finding')} ({f.get('source', '')})"
        for f in research.get("key_findings", [])[:5]
    ]

    prompt = CAM_EXECUTIVE_SUMMARY_PROMPT.format(
        company_name=state.get("company_name", ""),
        sector=state.get("sector", ""),
        recommendation=rec.get("recommendation", "REJECT"),
        loan_amount=rec.get("suggested_loan_amount", 0),
        interest_rate=rec.get("suggested_interest_rate", "N/A"),
        weighted_total=round(scores.get("weighted_total", 0), 1),
        five_cs_summary=five_cs_summary,
        red_flags="\n".join(f"  - {f}" for f in red_flags[:10]) or "  None",
        research_findings="\n".join(f"  - {f}" for f in key_findings) or "  None",
    )
    try:
        genai.configure(api_key=os.environ["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-1.5-pro')
        
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=2048,
                temperature=0.7,
            )
        )
        return response.text
    except Exception:
        return (
            f"This Credit Appraisal Memo summarises the credit assessment for "
            f"{state.get('company_name', 'the applicant')}. "
            f"The final recommendation is {rec.get('recommendation', 'PENDING')} "
            f"based on a Five Cs weighted score of {round(scores.get('weighted_total', 0), 1)}/100. "
            f"{rec.get('decision_reason', '')}"
        )


# ─── Document Styling ─────────────────────────────────────────────────────────

def _set_heading_style(paragraph, level: int = 1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)  # Teal
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x00, 0x99, 0xFF)  # Blue
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x4A, 0x60, 0x70)  # Muted


def _format_data_quality_section(state: Dict[str, Any]) -> str:
    """
    Generate Data Quality & Extraction Confidence section content.
    
    Returns formatted text describing document quality and confidence metrics.
    """
    extracted = state.get("extracted_financials", {})
    doc_quality = extracted.get("document_quality_summary", {})
    confidence_filtering = extracted.get("confidence_filtering", {})
    
    lines = []
    
    # Overall document quality
    total_docs = doc_quality.get("total_documents", 0)
    high_conf = doc_quality.get("high_confidence_count", 0)
    moderate_conf = doc_quality.get("moderate_confidence_count", 0)
    low_conf = doc_quality.get("low_confidence_count", 0)
    manual_review = doc_quality.get("manual_review_required", 0)
    
    lines.append(f"Total Documents Processed: {total_docs}")
    lines.append(f"  • High Confidence (≥70%): {high_conf} documents")
    lines.append(f"  • Moderate Confidence (50-70%): {moderate_conf} documents")
    lines.append(f"  • Low Confidence (<50%): {low_conf} documents")
    
    if manual_review > 0:
        lines.append(f"\n⚠ ALERT: {manual_review} document(s) require manual verification")
    
    # Document-level details
    if doc_quality.get("document_details"):
        lines.append("\nDocument Quality Breakdown:")
        for doc_detail in doc_quality["document_details"]:
            status_icon = "✓" if doc_detail["status"] == "ACCEPTED" else "⚠" if "WARNING" in doc_detail["status"] else "✗"
            lines.append(
                f"  {status_icon} {doc_detail['file_name']}: "
                f"{doc_detail['confidence']:.1%} ({doc_detail['reliability']})"
            )
    
    # Metric-level filtering
    if confidence_filtering:
        total_metrics = confidence_filtering.get("total_metrics", 0)
        accepted = confidence_filtering.get("accepted_metrics", 0)
        flagged = confidence_filtering.get("flagged_metrics", 0)
        
        lines.append(f"\nFinancial Metric Confidence Filtering:")
        lines.append(f"  • Total Metrics Extracted: {total_metrics}")
        lines.append(f"  • High Confidence Metrics: {accepted}")
        lines.append(f"  • Low Confidence Metrics (Flagged): {flagged}")
        
        if flagged > 0:
            lines.append("\n⚠ Low Confidence Metrics Requiring Verification:")
            for flagged_item in confidence_filtering.get("flagged_list", [])[:10]:
                lines.append(
                    f"  • {flagged_item['metric']}: {flagged_item['confidence']:.1%} confidence - "
                    f"Value: {flagged_item.get('value', 'N/A')}"
                )
    
    # Overall assessment
    if low_conf > 0 or (confidence_filtering and confidence_filtering.get("flagged_metrics", 0) > 0):
        lines.append("\n" + "="*60)
        lines.append("DATA QUALITY WARNING:")
        lines.append(
            "Some financial data was extracted with low confidence. "
            "It is recommended that the credit officer manually verify the flagged metrics "
            "against the original documents before making a final credit decision."
        )
        lines.append("="*60)
    else:
        lines.append("\n✓ All documents processed with acceptable confidence levels.")
    
    return "\n".join(lines)


def _add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        p.runs[0].font.color.rgb = (
            RGBColor(0x00, 0xD4, 0xAA) if level == 1
            else RGBColor(0x00, 0x99, 0xFF)
        )
    return p


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_row[i].text = h
        hdr_row[i].paragraphs[0].runs[0].bold = True if hdr_row[i].paragraphs[0].runs else None

    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = str(cell_text)
    doc.add_paragraph()  # Spacing after table


def _severity_emoji(severity: str) -> str:
    return {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(severity.upper(), "⚪")


def _rec_label(rec: str) -> str:
    labels = {
        "APPROVE": "✅ APPROVED",
        "CONDITIONAL APPROVE": "⚠ CONDITIONAL APPROVAL",
        "REJECT": "❌ REJECTED",
    }
    return labels.get(rec, rec)


# ─── Section Builders ─────────────────────────────────────────────────────────

def _build_cover_page(doc: Document, state: Dict[str, Any], rec: Dict[str, Any]):
    doc.add_paragraph()
    title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(state.get("company_name", "").upper())
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()
    meta = [
        f"Job ID: {state.get('job_id', 'N/A')}",
        f"Sector: {state.get('sector', 'N/A')}",
        f"Date: {datetime.now().strftime('%d %B %Y')}",
        f"Loan Requested: Rs. {state.get('loan_amount_requested', 0):.2f} Cr",
        f"Prepared by: Intelli-Credit AI Engine",
    ]
    for line in meta:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Decision stamp
    recommendation = rec.get("recommendation", "REJECT")
    stamp_text = _rec_label(recommendation)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(stamp_text)
    run.bold = True
    run.font.size = Pt(18)
    if "APPROVE" in recommendation and "REJECT" not in recommendation:
        run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)
    else:
        run.font.color.rgb = RGBColor(0xEF, 0x47, 0x6F)

    if rec.get("suggested_loan_amount"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(
            f"Rs. {rec.get('suggested_loan_amount', 0):.2f} Cr  |  "
            f"{rec.get('suggested_interest_rate', 'N/A')}"
        )
        r2.bold = True
        r2.font.size = Pt(14)

    doc.add_page_break()


def _build_executive_summary(doc: Document, summary_text: str):
    _add_section_heading(doc, "1. Executive Summary")
    doc.add_paragraph(summary_text)
    doc.add_paragraph()


def _build_company_background(doc: Document, extracted: Dict[str, Any]):
    _add_section_heading(doc, "2. Company Background")
    name = extracted.get("company_name", "N/A")
    promoters = extracted.get("promoters", [])
    doc.add_paragraph(f"Company Name: {name}")
    doc.add_paragraph(f"Promoters / Key Management: {', '.join(promoters) or 'Not identified'}")
    collateral = extracted.get("collateral", {})
    doc.add_paragraph(
        f"Collateral Offered: {collateral.get('type', 'N/A')} — "
        f"Estimated Value: Rs. {collateral.get('estimated_value', 0):.2f} Cr"
    )
    doc.add_paragraph()


def _build_financial_analysis(doc: Document, extracted: Dict[str, Any], loan_amount: float):
    _add_section_heading(doc, "3. Financial Analysis")

    fin = extracted.get("financials", {})
    revenue = fin.get("revenue_3yr", [0, 0, 0])
    ebitda = fin.get("ebitda_3yr", [0, 0, 0])
    pat = fin.get("pat_3yr", [0, 0, 0])

    years = ["FY22", "FY23", "FY24"]

    # Pad to 3 years
    while len(revenue) < 3:
        revenue = [0] + revenue
    while len(ebitda) < 3:
        ebitda = [0] + ebitda
    while len(pat) < 3:
        pat = [0] + pat

    _add_section_heading(doc, "3.1 3-Year Financial Performance (Rs. Crore)", level=2)
    headers = ["Metric"] + years
    rows = [
        ["Revenue (Net Sales)"] + [f"{v:.2f}" for v in revenue[-3:]],
        ["EBITDA"] + [f"{v:.2f}" for v in ebitda[-3:]],
        ["PAT (Net Profit)"] + [f"{v:.2f}" for v in pat[-3:]],
    ]
    _add_table(doc, headers, rows)

    _add_section_heading(doc, "3.2 Key Financial Ratios", level=2)
    ratio_headers = ["Ratio", "Value", "Benchmark", "Assessment"]
    dscr = fin.get("dscr", 0)
    d2e = fin.get("debt_to_equity", 0)
    ratio_rows = [
        ["DSCR", f"{dscr:.2f}x", "> 1.25x", "✅ Adequate" if dscr >= 1.25 else "❌ Weak"],
        ["Debt-to-Equity", f"{d2e:.2f}x", "< 3x", "✅ Acceptable" if d2e <= 3 else "❌ Over-leveraged"],
        ["Net Worth", f"Rs. {fin.get('net_worth', 0):.2f} Cr", "—", "—"],
        ["Total Debt", f"Rs. {fin.get('total_debt', 0):.2f} Cr", "—", "—"],
        ["CFO", f"Rs. {fin.get('cash_flow_from_operations', 0):.2f} Cr", "> 0", "✅ Positive" if fin.get("cash_flow_from_operations", 0) > 0 else "❌ Negative"],
    ]
    _add_table(doc, ratio_headers, ratio_rows)


def _build_data_quality_section(doc: Document, state: Dict[str, Any]):
    """
    Build Data Quality & Extraction Confidence section in CAM.
    
    Shows document quality metrics and confidence levels.
    """
    _add_section_heading(doc, "3.3 Data Quality & Extraction Confidence")
    
    quality_text = _format_data_quality_section(state)
    
    # Split into paragraphs and add with appropriate formatting
    for line in quality_text.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line)
            
            # Apply warning formatting for alert lines
            if "ALERT" in line or "WARNING" in line:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)  # Orange
                    run.bold = True
            elif "✓" in line or "All documents processed" in line:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x99, 0x00)  # Green
    
    doc.add_paragraph()


def _build_five_cs(doc: Document, scores: Dict[str, Any]):
    _add_section_heading(doc, "4. Five Cs of Credit Assessment")

    cs_labels = {
        "character": "Character (Promoter Integrity & Track Record)",
        "capacity": "Capacity (Ability to Repay)",
        "capital": "Capital (Financial Strength)",
        "collateral": "Collateral (Security Cover)",
        "conditions": "Conditions (Market & Sector Outlook)",
    }
    weights = {"character": "25%", "capacity": "30%", "capital": "20%", "collateral": "15%", "conditions": "10%"}

    for c_key, c_label in cs_labels.items():
        c_data = scores.get(c_key, {})
        score = c_data.get("score", 0)
        reasons = c_data.get("reasons", [])
        weight = weights.get(c_key, "")

        _add_section_heading(doc, f"4.{list(cs_labels.keys()).index(c_key)+1}  {c_label}  [{weight} weight]", level=2)
        p = doc.add_paragraph()
        r = p.add_run(f"Score: {score}/100")
        r.bold = True
        r.font.size = Pt(13)

        for reason in reasons:
            doc.add_paragraph(f"• {reason}", style="List Bullet")
        doc.add_paragraph()


def _build_score_breakdown_table(doc: Document, scores: Dict[str, Any]):
    _add_section_heading(doc, "5. Score Breakdown", level=1)
    headers = ["Five C", "Score (/100)", "Weight", "Weighted Score"]
    cs = ["character", "capacity", "capital", "collateral", "conditions"]
    rows = []
    for c in cs:
        c_data = scores.get(c, {})
        s = float(c_data.get("score", 0))
        w = float(c_data.get("weight", 0))
        rows.append([c.capitalize(), f"{s:.0f}", f"{w*100:.0f}%", f"{s*w:.1f}"])
    rows.append(["TOTAL", "", "", f"{scores.get('weighted_total', 0):.1f}"])
    _add_table(doc, headers, rows)


def _build_risk_flags(doc: Document, extracted: Dict[str, Any]):
    _add_section_heading(doc, "6. Risk Flags")
    flags = extracted.get("red_flags", [])
    if not flags:
        doc.add_paragraph("No significant risk flags identified from document analysis.")
        return
    for i, flag in enumerate(flags, 1):
        doc.add_paragraph(f"{i}. {flag}", style="List Number")
    doc.add_paragraph()


def _build_research_findings(doc: Document, research: Dict[str, Any]):
    _add_section_heading(doc, "7. Research & Due Diligence Findings")

    doc.add_paragraph(f"Litigation Risk: {research.get('litigation_risk', 'N/A')}")
    doc.add_paragraph(f"Promoter Integrity Score: {research.get('promoter_integrity_score', 'N/A')}/100")
    doc.add_paragraph(f"Sector Outlook: {research.get('sector_outlook', 'N/A')}")

    headwinds = research.get("sector_headwinds", [])
    if headwinds:
        _add_section_heading(doc, "Sector Headwinds", level=2)
        for hw in headwinds:
            doc.add_paragraph(f"• {hw}", style="List Bullet")

    findings = research.get("key_findings", [])
    if findings:
        _add_section_heading(doc, "Key Findings", level=2)
        headers = ["#", "Finding", "Severity", "Source", "Date"]
        rows = [
            [
                str(i),
                f.get("finding", "")[:80],
                f"{_severity_emoji(f.get('severity','LOW'))} {f.get('severity','LOW')}",
                f.get("source", "N/A"),
                f.get("date", "N/A"),
            ]
            for i, f in enumerate(findings, 1)
        ]
        _add_table(doc, headers, rows)

    impact = research.get("recommendation_impact", "")
    if impact:
        doc.add_paragraph(f"Impact on Decision: {impact}")
    doc.add_paragraph()


def _build_recommendation(doc: Document, state: Dict[str, Any], rec: Dict[str, Any], scores: Dict[str, Any]):
    _add_section_heading(doc, "8. Final Recommendation")

    recommendation = rec.get("recommendation", "REJECT")
    p = doc.add_paragraph()
    r = p.add_run(_rec_label(recommendation))
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(f"Suggested Loan Amount: Rs. {rec.get('suggested_loan_amount', 0):.2f} Crore")
    doc.add_paragraph(f"Suggested Interest Rate: {rec.get('suggested_interest_rate', 'N/A')}")
    doc.add_paragraph(f"Five Cs Weighted Score: {scores.get('weighted_total', 0):.1f}/100")
    doc.add_paragraph()

    doc.add_paragraph(rec.get("decision_reason", ""))
    doc.add_paragraph()

    overrides = rec.get("overriding_factors", [])
    if overrides:
        _add_section_heading(doc, "Overriding Factors / Conditions", level=2)
        for factor in overrides:
            doc.add_paragraph(f"• {factor}", style="List Bullet")
    doc.add_paragraph()

    # Turnaround time
    started = state.get("started_at", "")
    completed = state.get("completed_at", datetime.now().isoformat())
    if started:
        try:
            from datetime import timezone
            t0 = datetime.fromisoformat(started)
            t1 = datetime.fromisoformat(completed)
            delta = t1 - t0
            mins, secs = divmod(int(delta.total_seconds()), 60)
            doc.add_paragraph(
                f"Appraisal completed in: {mins}m {secs}s  "
                f"(Industry average: 10–15 business days)"
            )
        except Exception:
            pass


def _build_appendix(doc: Document, state: Dict[str, Any]):
    doc.add_page_break()
    _add_section_heading(doc, "Appendix — Raw Data")

    extracted = state.get("extracted_financials", {})
    research = state.get("research_findings", {})

    _add_section_heading(doc, "A1. Extracted Financials (raw)", level=2)
    doc.add_paragraph(json.dumps(extracted, indent=2))

    _add_section_heading(doc, "A2. Research Findings (raw)", level=2)
    doc.add_paragraph(json.dumps(research, indent=2))

    _add_section_heading(doc, "A3. Arbitration Record", level=2)
    arb = state.get("arbitration_result", {})
    doc.add_paragraph(json.dumps(arb, indent=2) if arb else "No arbitration was required.")

    searches = research.get("_search_queries", [])
    if searches:
        _add_section_heading(doc, "A4. Web Search Queries Executed", level=2)
        for s in searches:
            doc.add_paragraph(f"[{s.get('label')}] {s.get('query')}", style="List Bullet")


# ─── SWOT Analysis ────────────────────────────────────────────────────────────

async def generate_swot_analysis(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Generate SWOT analysis using Gemini with comprehensive state data.
    Falls back to rule-based SWOT if Gemini fails.
    """
    try:
        # Extract all relevant data
        company_name = state.get("company_name", "Company")
        sector = state.get("sector", "Unknown")
        loan_amount = state.get("loan_amount_requested", 0) / 10_000_000  # Convert to Cr
        
        extracted = state.get("extracted_financials", {})
        financials = extracted.get("financials", {})
        research = state.get("research_findings", {})
        scores = state.get("five_cs_scores", {})
        red_flags = extracted.get("red_flags", [])
        
        # Financial metrics
        revenue_3yr = financials.get("revenue_3yr_cagr", 0)
        ebitda_margin = financials.get("ebitda_margin", 0)
        pat_margin = financials.get("pat_margin", 0)
        dscr = financials.get("dscr", 0)
        debt_to_equity = financials.get("debt_to_equity", 0)
        current_ratio = financials.get("current_ratio", 0)
        interest_coverage = financials.get("interest_coverage", 0)
        
        # Research insights
        litigation_risk = research.get("litigation_risk", "N/A")
        promoter_score = research.get("promoter_integrity_score", 0)
        news_sentiment = research.get("news_sentiment", "Neutral")
        mca_status = research.get("mca_status", "N/A")
        
        # Sector analysis
        sector_analysis = research.get("sector_analysis", {})
        sector_status = sector_analysis.get("sector_status", "N/A")
        risk_score = sector_analysis.get("risk_score", 0)
        recommendation = sector_analysis.get("recommendation", "N/A")
        
        # Qualitative scores
        qualitative_inputs = state.get("qualitative_inputs", {})
        factory_score = qualitative_inputs.get("factory_visit", {}).get("overall_score", 0)
        management_score = qualitative_inputs.get("management_interview", {}).get("overall_score", 0)
        
        # Collateral
        collateral = scores.get("collateral", {})
        coverage_ratio = collateral.get("coverage_ratio", 0)
        marketability = collateral.get("marketability", "N/A")
        
        # Top 5 red flags
        top_5_flags = [f[:100] for f in red_flags[:5]]
        flags_str = "; ".join(top_5_flags) if top_5_flags else "None"
        
        # Build comprehensive Gemini prompt
        prompt = f"""Perform a SWOT analysis for {company_name} ({sector}) applying for ₹{loan_amount:.2f} Cr loan.

FINANCIALS: Revenue 3yr CAGR={revenue_3yr:.1f}%, EBITDA%={ebitda_margin:.1f}%, PAT%={pat_margin:.1f}%, 
DSCR={dscr:.2f}x, D/E={debt_to_equity:.2f}x, Current Ratio={current_ratio:.2f}x, ICR={interest_coverage:.2f}x.

RESEARCH: Litigation Risk={litigation_risk}, Promoter Score={promoter_score}/100,
News Sentiment={news_sentiment}, MCA Status={mca_status}.

RED FLAGS ({len(red_flags)} total): {flags_str}

SECTOR: Status={sector_status}, Risk Score={risk_score}/100, Outlook={recommendation}.

QUALITATIVE: Factory Score={factory_score}/100, Mgmt Score={management_score}/100.

COLLATERAL: Coverage={coverage_ratio:.2f}x, Marketability={marketability}.

Generate a bank-grade SWOT analysis with:
- 4-6 STRENGTHS: Positive financial ratios, competitive advantages, strong metrics
- 4-6 WEAKNESSES: Red flags, poor ratios, operational concerns
- 3-5 OPPORTUNITIES: Growth potential, sector trends, expansion possibilities
- 3-5 THREATS: Market risks, regulatory changes, competitive pressures

Use actual numbers from the data provided. Be specific and professional.

Return ONLY valid JSON in this exact format:
{{
  "strengths": ["strength 1", "strength 2", ...],
  "weaknesses": ["weakness 1", "weakness 2", ...],
  "opportunities": ["opportunity 1", "opportunity 2", ...],
  "threats": ["threat 1", "threat 2", ...],
  "overall_assessment": "2-3 sentence summary of SWOT implications",
  "key_consideration": "Single most critical factor for credit decision"
}}"""
        
        # Call Gemini
        genai.configure(api_key=os.environ["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-1.5-pro')
        
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=2048,
                temperature=0.7,
            )
        )
        
        # Parse JSON response
        response_text = response.text.strip()
        # Extract JSON from response (may be wrapped in markdown)
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        
        swot = json.loads(response_text)
        
        # Validate structure
        required_keys = ["strengths", "weaknesses", "opportunities", "threats", "overall_assessment", "key_consideration"]
        if not all(key in swot for key in required_keys):
            raise ValueError("Invalid SWOT structure from Gemini")
        
        return swot
        
    except Exception as e:
        # Fallback to rule-based SWOT
        print(f"Gemini SWOT generation failed: {e}. Using rule-based fallback.")
        return _generate_fallback_swot(state)


def _generate_fallback_swot(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Rule-based SWOT generation fallback when Gemini fails.
    Based on quantitative thresholds from state data.
    """
    extracted = state.get("extracted_financials", {})
    financials = extracted.get("financials", {})
    research = state.get("research_findings", {})
    scores = state.get("five_cs_scores", {})
    red_flags = extracted.get("red_flags", [])
    
    strengths = []
    weaknesses = []
    opportunities = []
    threats = []
    
    # ── STRENGTHS (6 rules) ───────────────────────────────────────────────────
    
    # 1. Strong DSCR
    dscr = financials.get("dscr", 0)
    if dscr >= 2.0:
        strengths.append(f"Excellent debt servicing capability with DSCR of {dscr:.2f}x (well above 1.5x minimum)")
    elif dscr >= 1.5:
        strengths.append(f"Adequate debt servicing with DSCR of {dscr:.2f}x")
    
    # 2. Healthy profitability
    ebitda_margin = financials.get("ebitda_margin", 0)
    if ebitda_margin >= 15:
        strengths.append(f"Strong profitability with EBITDA margin of {ebitda_margin:.1f}%")
    
    # 3. Low leverage
    debt_to_equity = financials.get("debt_to_equity", 0)
    if debt_to_equity <= 1.5:
        strengths.append(f"Conservative leverage at D/E of {debt_to_equity:.2f}x")
    
    # 4. Good liquidity
    current_ratio = financials.get("current_ratio", 0)
    if current_ratio >= 1.5:
        strengths.append(f"Strong liquidity position with current ratio of {current_ratio:.2f}x")
    
    # 5. Promoter integrity
    promoter_score = research.get("promoter_integrity_score", 0)
    if promoter_score >= 75:
        strengths.append(f"High promoter integrity score of {promoter_score}/100 indicating trustworthy management")
    
    # 6. Revenue growth
    revenue_cagr = financials.get("revenue_3yr_cagr", 0)
    if revenue_cagr >= 15:
        strengths.append(f"Robust revenue growth at {revenue_cagr:.1f}% 3-year CAGR")
    
    # Default strength if none qualify
    if not strengths:
        strengths.append("Company has operational history and established business presence")
    
    # ── WEAKNESSES (4 rules) ──────────────────────────────────────────────────
    
    # 1. High red flag count
    if len(red_flags) >= 5:
        weaknesses.append(f"{len(red_flags)} red flags identified including regulatory and financial concerns")
    
    # 2. Poor DSCR
    if dscr < 1.25 and dscr > 0:
        weaknesses.append(f"Weak debt servicing ability with DSCR of {dscr:.2f}x (below 1.25x threshold)")
    
    # 3. High leverage
    if debt_to_equity > 3.0:
        weaknesses.append(f"High leverage risk with D/E ratio of {debt_to_equity:.2f}x")
    
    # 4. Litigation risk
    litigation_risk = research.get("litigation_risk", "LOW")
    if litigation_risk in ["HIGH", "CRITICAL"]:
        weaknesses.append(f"Significant litigation exposure rated as {litigation_risk}")
    
    # Default weakness
    if not weaknesses:
        weaknesses.append("Limited operational track record requires continued monitoring")
    
    # ── OPPORTUNITIES (2 rules) ───────────────────────────────────────────────
    
    # 1. Sector outlook
    sector_analysis = research.get("sector_analysis", {})
    sector_status = sector_analysis.get("sector_status", "")
    if "GROWING" in sector_status.upper() or "STABLE" in sector_status.upper():
        opportunities.append(f"Favorable sector dynamics with {sector_status.lower()} market conditions")
    
    # 2. Expansion capacity
    capacity_score = scores.get("capacity", {}).get("score", 0)
    if capacity_score >= 70:
        opportunities.append("Strong operational capacity allows for business expansion and scaling")
    
    # Default opportunities
    if not opportunities:
        opportunities.append("Market expansion potential in existing product/service lines")
        opportunities.append("Operational efficiency improvements can enhance margins")
    
    # ── THREATS (2 rules) ─────────────────────────────────────────────────────
    
    # 1. Sector headwinds
    headwinds = research.get("sector_headwinds", [])
    if headwinds:
        threats.append(f"Sector faces headwinds including: {', '.join(headwinds[:2])}")
    
    # 2. Working capital stress
    working_capital_days = financials.get("working_capital_days", 0)
    if working_capital_days > 120:
        threats.append(f"Working capital cycle of {working_capital_days:.0f} days indicates cash flow pressure")
    
    # Default threats
    if not threats:
        threats.append("Competitive market dynamics and pricing pressures")
        threats.append("Regulatory changes in the sector may impact operations")
    
    # ── ASSESSMENT ────────────────────────────────────────────────────────────
    
    weighted_score = scores.get("weighted_total", 0)
    recommendation = state.get("final_recommendation", {}).get("recommendation", "PENDING")
    
    if weighted_score >= 70:
        overall_assessment = f"Overall strong credit profile with weighted score of {weighted_score:.1f}/100. Strengths significantly outweigh weaknesses, supporting {recommendation} recommendation."
    elif weighted_score >= 50:
        overall_assessment = f"Moderate credit profile with weighted score of {weighted_score:.1f}/100. Balance of strengths and weaknesses requires careful evaluation. Recommendation: {recommendation}."
    else:
        overall_assessment = f"Weak credit profile with weighted score of {weighted_score:.1f}/100. Weaknesses and threats dominate analysis. Recommendation: {recommendation}."
    
    # Key consideration
    if len(red_flags) >= 5:
        key_consideration = f"Address {len(red_flags)} identified red flags before credit approval"
    elif dscr < 1.25 and dscr > 0:
        key_consideration = "Debt servicing capability below minimum threshold requires enhancement"
    elif promoter_score < 50:
        key_consideration = "Low promoter integrity score requires detailed background verification"
    else:
        key_consideration = "Monitor financial performance and ensure adherence to loan covenants"
    
    return {
        "strengths": strengths[:6],  # Max 6
        "weaknesses": weaknesses[:6],  # Max 6
        "opportunities": opportunities[:5],  # Max 5
        "threats": threats[:5],  # Max 5
        "overall_assessment": overall_assessment,
        "key_consideration": key_consideration
    }


def _build_swot_section(doc: Document, swot: Dict[str, Any]):
    """Format SWOT analysis section for Word document."""
    _add_section_heading(doc, "5. SWOT ANALYSIS")
    
    # Overall Assessment
    doc.add_paragraph(swot.get("overall_assessment", ""), style="Body Text")
    doc.add_paragraph()
    
    # Strengths
    _add_section_heading(doc, "5.1 STRENGTHS ✓", level=2)
    for i, strength in enumerate(swot.get("strengths", []), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(strength)
    doc.add_paragraph()
    
    # Weaknesses
    _add_section_heading(doc, "5.2 WEAKNESSES ⚠", level=2)
    for i, weakness in enumerate(swot.get("weaknesses", []), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(weakness)
    doc.add_paragraph()
    
    # Opportunities
    _add_section_heading(doc, "5.3 OPPORTUNITIES ↗", level=2)
    for i, opportunity in enumerate(swot.get("opportunities", []), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(opportunity)
    doc.add_paragraph()
    
    # Threats
    _add_section_heading(doc, "5.4 THREATS ↘", level=2)
    for i, threat in enumerate(swot.get("threats", []), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(threat)
    doc.add_paragraph()
    
    # Key Consideration
    _add_section_heading(doc, "5.5 KEY CONSIDERATION", level=2)
    p = doc.add_paragraph()
    run = p.add_run(swot.get("key_consideration", ""))
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x69, 0x00)  # Orange-red for emphasis
    doc.add_paragraph()


# ─── Main Agent Function ──────────────────────────────────────────────────────

async def run_cam_generator(state: Dict[str, Any]) -> Dict[str, Any]:
    """LangGraph node: CAM Generator Agent."""
    logs: List[Dict[str, Any]] = []
    agent = "CAM"

    company_name = state.get("company_name", "Company")
    job_id = state.get("job_id", "unknown")
    scores = state.get("five_cs_scores", {})
    rec = state.get("final_recommendation", {})
    extracted = state.get("extracted_financials", {})
    research = state.get("research_findings", {})

    logs.append(_log(agent, f"Generating Credit Appraisal Memo for '{company_name}'..."))

    # ── Get executive summary from Gemini ─────────────────────────────────────
    logs.append(_log(agent, "Requesting Executive Summary from Gemini..."))
    exec_summary = _call_gemini_for_summary(state)
    logs.append(_log(agent, "Executive Summary generated."))

    # ── Generate SWOT Analysis ────────────────────────────────────────────────
    logs.append(_log(agent, "Generating SWOT Analysis..."))
    swot_analysis = await generate_swot_analysis(state)
    logs.append(_log(agent, f"SWOT Analysis complete: {len(swot_analysis.get('strengths', []))} strengths, "
                             f"{len(swot_analysis.get('weaknesses', []))} weaknesses identified"))

    # ── Build Word document ───────────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    logs.append(_log(agent, "Building CAM document sections..."))

    state_with_completed = {**state, "completed_at": state.get("completed_at") or datetime.now().isoformat()}

    _build_cover_page(doc, state_with_completed, rec)
    _build_executive_summary(doc, exec_summary)
    _build_company_background(doc, extracted)
    _build_financial_analysis(doc, extracted, state.get("loan_amount_requested", 0))
    _build_data_quality_section(doc, state)  # NEW: Data Quality & Extraction Confidence
    _build_swot_section(doc, swot_analysis)  # NEW: SWOT Analysis
    _build_five_cs(doc, scores)
    _build_score_breakdown_table(doc, scores)
    _build_risk_flags(doc, extracted)
    _build_research_findings(doc, research)
    _build_recommendation(doc, state_with_completed, rec, scores)
    _build_appendix(doc, state_with_completed)

    # ── Save document ─────────────────────────────────────────────────────────
    upload_dir = os.getenv("UPLOAD_DIR", "./uploads")
    output_dir = Path(upload_dir) / job_id
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_name = "".join(c for c in company_name if c.isalnum() or c in " _-").strip().replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cam_filename = f"CAM_{safe_name}_{timestamp}.docx"
    cam_path = str(output_dir / cam_filename)

    doc.save(cam_path)

    logs.append(_log(agent, f"CAM document saved: {cam_filename}", level="SUCCESS"))
    logs.append(_log(agent, "CAM Generator Agent DONE.", level="SUCCESS"))

    return {
        "cam_path": cam_path,
        "swot_analysis": swot_analysis,  # Include SWOT for Results page
        "completed_at": datetime.now().isoformat(),
        "status": "COMPLETED",
        "logs": logs,
        "agent_statuses": {**state.get("agent_statuses", {}), "cam_generator": "DONE"},
    }
